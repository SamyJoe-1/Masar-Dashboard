# core/export.py
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------- utils ----------
def _s(x):
    if x is None: return ""
    if isinstance(x, (list, tuple)):
        return ", ".join(_s(v) for v in x if v is not None)
    if isinstance(x, dict):
        # كُن واضح: لا نكتب dict مباشرة في الـ PDF/Docx
        return "; ".join(f"{k}: {_s(v)}" for k, v in x.items())
    return str(x)

def _one(x: str) -> str:
    return _s(x).replace("\r", " ").replace("\n", " ").strip()

# ----------- DOCX (CV) ----------
def export_docx(cv: dict) -> bytes:
    doc = Document()
    # Header
    h = doc.add_heading(_one(cv.get("name") or "Candidate Name"), level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(_one(cv.get("title") or "Professional Title")).bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{_one(cv.get('location') or '-')}" + " | " +
              f"{_one(cv.get('email') or '-')}" + " | " +
              f"{_one(cv.get('phone') or '-')}")
    links_line = " | ".join([x for x in [_one(cv.get("linkedin")), _one(cv.get("links"))] if x and x != "-"])
    if links_line:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(links_line)

    # Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(_one(cv.get("summary")) or "-")

    # Skills & Tools
    if cv.get("skills"):
        doc.add_heading("Core Competencies", level=1)
        doc.add_paragraph(", ".join(_one(s) for s in cv.get("skills")))
    if cv.get("tools"):
        doc.add_heading("Tools & Tech", level=1)
        doc.add_paragraph(", ".join(_one(t) for t in cv.get("tools")))

    # Experience
    doc.add_heading("Experience — JD-tailored Results", level=1)
    exp = (cv.get("experience") or [{}])[0]
    bullets = exp.get("bullets") or []
    if bullets:
        for b in bullets:
            doc.add_paragraph(_one(b), style="List Bullet")
    else:
        doc.add_paragraph("-")

    # Projects
    if cv.get("projects"):
        doc.add_heading("Projects", level=1)
        for pjt in cv["projects"]:
            doc.add_paragraph(_one(pjt), style="List Bullet")

    # Education
    if cv.get("education"):
        doc.add_heading("Education", level=1)
        for ed in cv["education"]:
            doc.add_paragraph(_one(ed), style="List Bullet")

    # Certs
    if cv.get("certificates"):
        doc.add_heading("Certifications", level=1)
        doc.add_paragraph(", ".join(_one(c) for c in cv["certificates"]))

    # Languages
    if cv.get("languages"):
        doc.add_heading("Languages", level=1)
        doc.add_paragraph(", ".join(_one(l) for l in cv["languages"]))

    # Achievements
    if cv.get("achievements"):
        doc.add_heading("Achievements", level=1)
        for a in cv["achievements"]:
            doc.add_paragraph(_one(a), style="List Bullet")

    bio = BytesIO()
    doc.save(bio); bio.seek(0)
    return bio.getvalue()

# ----------- PDF (CV) ----------
def export_pdf(cv: dict) -> bytes:
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x_margin, y = 2*cm, height - 2*cm

    def draw(text, size=10, bold=False, leading=14):
        nonlocal y
        text = _one(text)
        if y < 2*cm:
            c.showPage(); y = height - 2*cm
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        c.drawString(x_margin, y, text)
        y -= leading

    def section(title):
        draw(title, size=12, bold=True, leading=18)

    # Header
    draw(_one(cv.get("name") or "Candidate Name"), size=16, bold=True, leading=20)
    draw(_one(cv.get("title") or "Professional Title"), size=11)
    draw(f"{_one(cv.get('location') or '-') }  |  {_one(cv.get('email') or '-') }  |  {_one(cv.get('phone') or '-') }", size=9)
    links_line = " | ".join([x for x in [_one(cv.get("linkedin")), _one(cv.get("links"))] if x and x != "-"])
    if links_line: draw(links_line, size=9)

    # Summary
    section("Executive Summary")
    for line in _one(cv.get("summary")).split(". "):
        if line.strip(): draw(line.strip() + ".")

    # Skills & Tools
    if cv.get("skills"):
        section("Core Competencies")
        draw(", ".join(_one(s) for s in cv["skills"]))
    if cv.get("tools"):
        section("Tools & Tech")
        draw(", ".join(_one(t) for t in cv["tools"]))

    # Experience
    section("Experience — JD-tailored Results")
    bullets = (cv.get("experience") or [{}])[0].get("bullets") or []
    if bullets:
        for b in bullets:
            draw("• " + _one(b))
    else:
        draw("-")

    # Projects
    if cv.get("projects"):
        section("Projects")
        for p in cv["projects"]:
            draw("• " + _one(p))

    # Education
    if cv.get("education"):
        section("Education")
        for ed in cv["education"]:
            draw("• " + _one(ed))

    # Certs
    if cv.get("certificates"):
        section("Certifications")
        draw(", ".join(_one(c) for c in cv["certificates"]))

    # Languages
    if cv.get("languages"):
        section("Languages")
        draw(", ".join(_one(l) for l in cv["languages"]))

    # Achievements
    if cv.get("achievements"):
        section("Achievements")
        for a in cv["achievements"]:
            draw("• " + _one(a))

    c.showPage(); c.save()
    buf.seek(0)
    return buf.getvalue()

# ----------- Cover Letter (DOCX) ----------
def export_docx_cover_letter(cv: dict, letter: str) -> bytes:
    doc = Document()
    # Header
    h = doc.add_heading("Cover Letter", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(_one(cv.get("name") or "Candidate Name")).bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{_one(cv.get('email') or '-') } | {_one(cv.get('phone') or '-') } | {_one(cv.get('linkedin') or '-') }")

    doc.add_paragraph("")  # spacer
    for para in (_s(letter).split("\n") or []):
        if para.strip():
            doc.add_paragraph(para.strip())
    bio = BytesIO(); doc.save(bio); bio.seek(0)
    return bio.getvalue()

# ----------- Cover Letter (PDF) ----------
def export_pdf_cover_letter(cv: dict, letter: str) -> bytes:
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x, y = 2*cm, height - 2*cm

    def draw(text, size=10, leading=14, bold=False):
        nonlocal y
        text = _one(text)
        if y < 2*cm:
            c.showPage(); y = height - 2*cm
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        c.drawString(x, y, text); y -= leading

    draw("Cover Letter", size=16, bold=True, leading=20)
    draw(_one(cv.get("name") or "Candidate Name"), size=12, bold=True)
    draw(f"{_one(cv.get('email') or '-') } | {_one(cv.get('phone') or '-') } | {_one(cv.get('linkedin') or '-') }", size=9)
    draw("")

    for para in (_s(letter).split("\n") or []):
        if para.strip():
            for line in para.strip().split(". "):
                if line.strip():
                    draw(line.strip() + ".")

    c.showPage(); c.save()
    buf.seek(0)
    return buf.getvalue()
